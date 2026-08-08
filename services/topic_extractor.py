import re
import json
import uuid
import bisect
from dataclasses import dataclass, field
from typing import Optional
from collections import Counter

from config import COHERE_API_KEY

# --- ADD THESE TO config.py ---
# TOPIC_AI_MODEL = "command-r-plus"
# TOPIC_AI_MAX_CHARS = 12000
from config import TOPIC_AI_MODEL, TOPIC_AI_MAX_CHARS


@dataclass
class Topic:
    id: str
    title: str
    level: int
    startPage: Optional[int] = None
    endPage: Optional[int] = None
    # internal-only, used to map chunks -> topics, not part of the public shape
    char_start: Optional[int] = field(default=None, repr=False)
    char_end: Optional[int] = field(default=None, repr=False)

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "title": self.title,
            "level": self.level,
            "startPage": self.startPage,
            "endPage": self.endPage,
        }


def _fill_end_pages_and_offsets(topics: list[Topic], total_pages: Optional[int] = None):
    for i, t in enumerate(topics):
        if i + 1 < len(topics):
            t.char_end = topics[i + 1].char_start
            if t.startPage is not None and topics[i + 1].startPage is not None:
                t.endPage = topics[i + 1].startPage
        else:
            t.char_end = None
            t.endPage = total_pages


# ────────────────────────── shared false-positive filters ──────────────────────────

_FALSE_POSITIVE_PATTERNS = [
    re.compile(r"^\d+$"),                                            # bare page number
    re.compile(r"^(page|pg\.?)\s*\d+(\s*(of|/)\s*\d+)?$", re.I),      # "Page 3", "Page 3 of 12"
    re.compile(r"^(figure|fig\.?|table|tbl\.?|chart|eq\.?|equation)\s*\d+", re.I),
    re.compile(r"^https?://\S+$", re.I),
    re.compile(r"^www\.\S+$", re.I),
    re.compile(r"^\d{1,2}[/.\-]\d{1,2}[/.\-]\d{2,4}$"),               # numeric date
    re.compile(r"^[A-Za-z]+\s+\d{1,2},?\s+\d{4}$"),                   # "January 4, 2024"
    re.compile(r"^copyright\b|^\(c\)\s*\d{4}|^all rights reserved", re.I),
    re.compile(r"^isbn\b", re.I),
]

_NUMBERED_HEADING_RE = re.compile(r"^(\d+(\.\d+)*)[.\)]?\s+\S")
_CHAPTER_SECTION_RE = re.compile(r"^(chapter|section|part|appendix)\s+([\dIVXLC]+)", re.I)


def _is_false_positive(text: str) -> bool:
    t = text.strip()
    if not t:
        return True
    return any(p.match(t) for p in _FALSE_POSITIVE_PATTERNS)


def _numbering_depth(text: str) -> Optional[int]:
    """Returns heading depth implied by a numeric/roman prefix, or None."""
    m = _NUMBERED_HEADING_RE.match(text.strip())
    if m:
        return m.group(1).count(".") + 1
    if _CHAPTER_SECTION_RE.match(text.strip()):
        return 1
    return None


# ────────────────────────── PDF ──────────────────────────

def _pdf_page_char_offsets(doc) -> dict[int, int]:
    """Mirrors '\\n'.join(page.get_text() for page in doc) from file_parser._parse_pdf"""
    offsets = {}
    cursor = 0
    for i, page in enumerate(doc):
        offsets[i + 1] = cursor
        cursor += len(page.get_text()) + 1
    return offsets


def _extract_topics_pdf_toc(doc) -> list[Topic]:
    toc = doc.get_toc(simple=True)  # [[level, title, page], ...], 1-indexed pages
    if not toc:
        return []
    page_offsets = _pdf_page_char_offsets(doc)
    topics = []
    for level, title, page in toc:
        topics.append(Topic(
            id=str(uuid.uuid4()),
            title=title.strip(),
            level=level,
            startPage=page,
            char_start=page_offsets.get(page, 0),
        ))
    _fill_end_pages_and_offsets(topics, total_pages=doc.page_count)
    return topics


def _rawdict_span_text(span: dict) -> str:
    """rawdict spans expose 'chars' (list of {'c': char, ...}) instead of 'text'."""
    if "text" in span:
        return span["text"]
    return "".join(c.get("c", "") for c in span.get("chars", []))


def _collect_pdf_lines(doc) -> list[dict]:
    """Flattens every non-empty line in the document into a metadata record,
    in reading order, using rawdict for richer font/layout info."""
    page_offsets = _pdf_page_char_offsets(doc)
    lines = []
    line_index = 0
    for i, page in enumerate(doc):
        page_num = i + 1
        page_width = page.rect.width
        page_height = page.rect.height
        cursor = 0
        prev_bottom = None
        for block in page.get_text("rawdict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                spans = line.get("spans", [])
                text = "".join(_rawdict_span_text(s) for s in spans)
                stripped = text.strip()
                if not stripped:
                    cursor += len(text) + 1
                    continue
                sizes = [s.get("size", 0) for s in spans]
                size = max(sizes) if sizes else 0
                bold = any((s.get("flags", 0) & (1 << 4)) or "bold" in s.get("font", "").lower()
                           for s in spans)
                font = spans[0].get("font", "") if spans else ""
                bbox = line.get("bbox", [0, 0, 0, 0])
                gap_before = (bbox[1] - prev_bottom) if prev_bottom is not None else None
                lines.append({
                    "line_index": line_index,
                    "page": page_num,
                    "text": stripped,
                    "size": size,
                    "bold": bold,
                    "font": font,
                    "bbox": bbox,
                    "gap_before": gap_before,
                    "char_start": page_offsets[page_num] + cursor,
                    "page_width": page_width,
                    "page_height": page_height,
                })
                prev_bottom = bbox[3]
                cursor += len(text) + 1
                line_index += 1
    return lines


def _pdf_body_stats(lines: list[dict]) -> tuple[float, float, str]:
    size_weight = Counter()
    font_weight = Counter()
    gaps = []
    for ln in lines:
        w = max(len(ln["text"]), 1)
        size_weight[round(ln["size"], 1)] += w
        font_weight[ln["font"]] += w
        if ln["gap_before"] is not None and ln["gap_before"] > 0:
            gaps.append(ln["gap_before"])
    body_size = size_weight.most_common(1)[0][0] if size_weight else 10.0
    body_font = font_weight.most_common(1)[0][0] if font_weight else ""
    gaps.sort()
    normal_gap = gaps[len(gaps) // 2] if gaps else 3.0
    return body_size, max(normal_gap, 1.0), body_font


def _detect_running_headers(lines: list[dict], page_count: int) -> set[str]:
    """Text that repeats near-verbatim across many pages is boilerplate
    (running headers/footers), not real headings."""
    normalize = lambda t: re.sub(r"\d+", "#", t.strip().lower())
    page_sets: dict[str, set[int]] = {}
    for ln in lines:
        key = normalize(ln["text"])
        page_sets.setdefault(key, set()).add(ln["page"])
    threshold = max(3, int(page_count * 0.15))
    return {k for k, pages in page_sets.items() if len(pages) >= threshold}


def _score_pdf_line(ln: dict, body_size: float, normal_gap: float, body_font: str) -> float:
    text = ln["text"]
    if _is_false_positive(text):
        return -100.0
    if len(text) > 160:
        return -100.0

    score = 0.0
    ratio = ln["size"] / body_size if body_size else 1.0
    if ratio > 1.5:
        score += 3.0
    elif ratio > 1.2:
        score += 2.0
    elif ratio > 1.05:
        score += 1.0

    if ln["bold"]:
        score += 1.5

    if ln["font"] and ln["font"] != body_font:
        score += 0.8

    if ln["gap_before"] is not None:
        if ln["gap_before"] > normal_gap * 2.5:
            score += 2.0
        elif ln["gap_before"] > normal_gap * 1.6:
            score += 1.0

    bbox = ln["bbox"]
    line_width = bbox[2] - bbox[0]
    center_x = (bbox[0] + bbox[2]) / 2
    page_center = ln["page_width"] / 2
    if ln["page_width"] and abs(center_x - page_center) < 0.08 * ln["page_width"] \
            and line_width < 0.75 * ln["page_width"]:
        score += 1.0

    depth = _numbering_depth(text)
    if depth is not None:
        score += 2.0

    if len(text) <= 70:
        score += 0.5
    elif len(text) > 140:
        score -= 2.0

    if text.isupper() and any(c.isalpha() for c in text):
        score += 1.2
    elif sum(1 for w in text.split() if w[:1].isupper()) >= max(1, len(text.split()) - 1):
        score += 0.4

    # de-prioritize the extreme top/bottom margin band (likely header/footer zone)
    if ln["page_height"]:
        y_frac = bbox[1] / ln["page_height"]
        if y_frac < 0.04 or y_frac > 0.96:
            score -= 1.5

    return score


HEADING_MIN_SCORE = 3.5


def _extract_topics_pdf_heuristic(doc) -> list[Topic]:
    lines = _collect_pdf_lines(doc)
    if not lines:
        return []

    body_size, normal_gap, body_font = _pdf_body_stats(lines)
    boilerplate = _detect_running_headers(lines, doc.page_count)
    normalize = lambda t: re.sub(r"\d+", "#", t.strip().lower())

    candidates = []
    for ln in lines:
        if normalize(ln["text"]) in boilerplate:
            continue
        score = _score_pdf_line(ln, body_size, normal_gap, body_font)
        if score >= HEADING_MIN_SCORE:
            candidates.append({**ln, "score": score})

    if not candidates:
        return []

    # merge adjacent lines (in original reading order) into single headings
    merged = []
    for c in candidates:
        if merged:
            prev = merged[-1]
            same_page = prev["page"] == c["page"]
            adjacent = c["line_index"] == prev["last_line_index"] + 1
            similar_size = abs(prev["size"] - c["size"]) < 0.6
            similar_bold = prev["bold"] == c["bold"]
            short_enough = len(prev["text"]) + len(c["text"]) < 160
            small_gap = c["gap_before"] is not None and c["gap_before"] < normal_gap * 2.0
            if same_page and adjacent and similar_size and similar_bold and short_enough and small_gap:
                prev["text"] = f'{prev["text"]} {c["text"]}'
                prev["last_line_index"] = c["line_index"]
                prev["score"] = max(prev["score"], c["score"])
                continue
        c["last_line_index"] = c["line_index"]
        merged.append(dict(c))

    # dynamic level ranking by font size, refined by explicit numbering depth
    unique_sizes = sorted({round(m["size"], 1) for m in merged}, reverse=True)
    size_rank = {sz: i + 1 for i, sz in enumerate(unique_sizes)}

    topics = []
    for m in merged:
        depth = _numbering_depth(m["text"])
        level = depth if depth is not None else size_rank.get(round(m["size"], 1), len(unique_sizes))
        level = max(1, min(level, 6))
        topics.append(Topic(
            id=str(uuid.uuid4()), title=m["text"], level=level,
            startPage=m["page"], char_start=m["char_start"],
        ))

    _fill_end_pages_and_offsets(topics, total_pages=doc.page_count)
    return topics


# ────────────────────────── DOCX ──────────────────────────

_DOCX_HEADING_STYLE_RE = re.compile(r"^(Heading)\s*(\d*)$", re.IGNORECASE)
_DOCX_TITLE_STYLE_RE = re.compile(r"^(Title|Subtitle)$", re.IGNORECASE)


def _docx_has_numbering(paragraph) -> bool:
    from docx.oxml.ns import qn
    pPr = paragraph._p.find(qn("w:pPr"))
    if pPr is None:
        return False
    return pPr.find(qn("w:numPr")) is not None


def _docx_outline_level(paragraph) -> Optional[int]:
    lvl = paragraph.paragraph_format.outline_level
    return lvl if lvl is not None else None


def _docx_run_font_size_pt(paragraph) -> Optional[float]:
    for r in paragraph.runs:
        if r.text.strip() and r.font.size is not None:
            return r.font.size.pt
    return None


def _docx_space_before_pt(paragraph) -> float:
    sb = paragraph.paragraph_format.space_before
    return sb.pt if sb is not None else 0.0


def _score_docx_paragraph(p, text: str, body_size: float) -> tuple[float, Optional[int]]:
    """Returns (score, forced_level). forced_level overrides font-based ranking
    when a strong structural signal (style/outline) is present."""
    if _is_false_positive(text) or len(text) > 160:
        return -100.0, None

    style_name = (p.style.name or "") if p.style else ""
    m = _DOCX_HEADING_STYLE_RE.match(style_name)
    if m:
        level = int(m.group(2)) if m.group(2) else 1
        return 10.0, level
    if _DOCX_TITLE_STYLE_RE.match(style_name):
        return 10.0, 1

    score = 0.0
    forced_level = None

    outline = _docx_outline_level(p)
    if outline is not None and outline < 9:
        score += 4.0
        forced_level = outline + 1

    depth = _numbering_depth(text)
    if depth is not None:
        score += 2.0
        forced_level = forced_level or depth

    runs = [r for r in p.runs if r.text.strip()]
    all_bold = bool(runs) and all(r.bold for r in runs)
    if all_bold and len(text) < 100:
        score += 1.5

    size = _docx_run_font_size_pt(p)
    if size and body_size:
        ratio = size / body_size
        if ratio > 1.5:
            score += 3.0
        elif ratio > 1.2:
            score += 2.0
        elif ratio > 1.05:
            score += 1.0

    if _docx_space_before_pt(p) > 12:
        score += 1.0

    if text.isupper() and any(c.isalpha() for c in text):
        score += 1.0

    if len(text) <= 70:
        score += 0.5
    elif len(text) > 140:
        score -= 2.0

    return score, forced_level


def _extract_topics_docx(path: str) -> list[Topic]:
    from docx import Document
    doc = Document(path)

    # body font size = most common run size across non-empty paragraphs
    size_weight = Counter()
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        sz = _docx_run_font_size_pt(p)
        if sz:
            size_weight[round(sz, 1)] += len(text)
    body_size = size_weight.most_common(1)[0][0] if size_weight else 11.0

    candidates = []
    cursor = 0
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        score, forced_level = _score_docx_paragraph(p, text, body_size)
        if score >= HEADING_MIN_SCORE:
            sz = _docx_run_font_size_pt(p) or body_size
            candidates.append({
                "text": text, "char_start": cursor,
                "forced_level": forced_level, "size": sz,
            })
        cursor += len(text) + 1

    if not candidates:
        return []

    unique_sizes = sorted({round(c["size"], 1) for c in candidates}, reverse=True)
    size_rank = {sz: i + 1 for i, sz in enumerate(unique_sizes)}

    topics = []
    for c in candidates:
        level = c["forced_level"] if c["forced_level"] is not None else size_rank.get(round(c["size"], 1), 1)
        level = max(1, min(level, 6))
        topics.append(Topic(id=str(uuid.uuid4()), title=c["text"], level=level, char_start=c["char_start"]))

    _fill_end_pages_and_offsets(topics)  # no page concept in docx
    return topics


# ────────────────────────── TXT / MD ──────────────────────────

_MD_ATX_RE = re.compile(r"^(#{1,6})\s+(.*?)#*\s*$")
_MD_SETEXT_RE = re.compile(r"^(=+|-{2,})\s*$")
_BOLD_LINE_RE = re.compile(r"^\*\*(.+?)\*\*$|^__(.+?)__$")
_FENCE_RE = re.compile(r"^```|^~~~")


def _extract_topics_text(text: str) -> list[Topic]:
    lines = text.split("\n")
    topics = []
    cursor = 0
    in_fence = False
    pending_title = None  # (text, char_start) awaiting a possible setext underline

    def flush_pending():
        # a held-back candidate line that wasn't followed by a setext underline
        # (===/---) was just plain body text, not a heading — discard it.
        nonlocal pending_title
        pending_title = None

    for line in lines:
        stripped = line.strip()

        if _FENCE_RE.match(stripped):
            flush_pending()
            in_fence = not in_fence
            cursor += len(line) + 1
            continue
        if in_fence:
            cursor += len(line) + 1
            continue

        setext_m = _MD_SETEXT_RE.match(stripped)
        if setext_m and pending_title:
            t, cs = pending_title
            level = 1 if stripped[0] == "=" else 2
            topics.append(Topic(id=str(uuid.uuid4()), title=t, level=level, char_start=cs))
            pending_title = None
            cursor += len(line) + 1
            continue

        flush_pending()

        if not stripped:
            cursor += len(line) + 1
            continue

        atx_m = _MD_ATX_RE.match(stripped)
        if atx_m:
            topics.append(Topic(
                id=str(uuid.uuid4()), title=atx_m.group(2).strip(),
                level=len(atx_m.group(1)), char_start=cursor,
            ))
            cursor += len(line) + 1
            continue

        depth = _numbering_depth(stripped) if not _is_false_positive(stripped) else None
        if depth is not None and 0 < len(stripped) < 100 and stripped[-1] not in ".,;:":
            topics.append(Topic(
                id=str(uuid.uuid4()), title=stripped, level=depth, char_start=cursor,
            ))
            cursor += len(line) + 1
            continue

        bold_m = _BOLD_LINE_RE.match(stripped)
        if bold_m and not _is_false_positive(stripped):
            title = (bold_m.group(1) or bold_m.group(2)).strip()
            topics.append(Topic(id=str(uuid.uuid4()), title=title, level=2, char_start=cursor))
            cursor += len(line) + 1
            continue

        if (stripped.isupper() and 3 < len(stripped) < 80 and any(c.isalpha() for c in stripped)
                and not _is_false_positive(stripped)):
            topics.append(Topic(id=str(uuid.uuid4()), title=stripped, level=1, char_start=cursor))
            cursor += len(line) + 1
            continue

        # a short, non-terminal-punctuated line might be a setext heading title;
        # hold it in case the next line is an underline
        if 0 < len(stripped) < 100 and stripped[-1] not in ".,;:" and not _is_false_positive(stripped):
            pending_title = (stripped, cursor)

        cursor += len(line) + 1

    flush_pending()
    topics.sort(key=lambda t: t.char_start)
    _fill_end_pages_and_offsets(topics)
    return topics


# ────────────────────────── AI fallback ──────────────────────────

def _extract_topics_ai(text: str) -> list[Topic]:
    """Only reached when no deterministic structure was found."""
    if not text.strip():
        return []

    sample = text[:TOPIC_AI_MAX_CHARS]
    prompt = (
        "You are analyzing a document to identify its topic/section structure. "
        "Return ONLY a JSON array (no prose, no markdown fences) of objects: "
        '{"title": string, "level": integer (1=top level), '
        '"snippet": string (first ~6 words of that section, copied exactly from the text)}. '
        "List topics in document order. If there's no clear structure, return one topic "
        "covering the whole document.\n\nTEXT:\n" + sample
    )

    try:
        import cohere
        client = cohere.Client(api_key=COHERE_API_KEY)
        resp = client.chat(message=prompt, model=TOPIC_AI_MODEL, temperature=0)
        raw = re.sub(r"^```(json)?|```$", "", resp.text.strip(), flags=re.MULTILINE).strip()
        parsed = json.loads(raw)
    except Exception:
        # last-resort: whole doc as one topic, never fail ingestion because of this
        return [Topic(id=str(uuid.uuid4()), title="Document", level=1, char_start=0, char_end=len(text))]

    topics = []
    for item in parsed:
        title = str(item.get("title", "")).strip()
        if not title:
            continue
        level = int(item.get("level") or 1)
        snippet = str(item.get("snippet", "")).strip()
        char_start = text.find(snippet) if snippet else -1
        topics.append(Topic(
            id=str(uuid.uuid4()), title=title, level=level,
            char_start=char_start if char_start != -1 else None,
        ))

    located = [t for t in topics if t.char_start is not None]
    located.sort(key=lambda t: t.char_start)
    _fill_end_pages_and_offsets(located)
    return located or topics


# ────────────────────────── dispatcher ──────────────────────────

def extract_topics(file_path: str, ext: str, text: str) -> list[Topic]:
    ext = ext.lower()
    topics: list[Topic] = []
    try:
        if ext == ".pdf":
            import fitz
            doc = fitz.open(file_path)
            topics = _extract_topics_pdf_toc(doc) or _extract_topics_pdf_heuristic(doc)
        elif ext in (".docx", ".doc"):
            topics = _extract_topics_docx(file_path)
        elif ext in (".txt", ".md", ".csv"):
            topics = _extract_topics_text(text)
    except Exception:
        topics = []

    if not topics:
        topics = _extract_topics_ai(text)

    return topics


def assign_topics_to_chunks(
    chunks_with_offsets: list[tuple[str, int]], topics: list[Topic]
) -> list[Optional[str]]:
    """Tags each chunk with the title of the topic it falls under.

    Runs in O((n + m) log m) via a single linear sweep over chunks sorted by
    offset against topics sorted by char_start, instead of the previous
    O(n * m) nested-loop scan.
    """
    n = len(chunks_with_offsets)
    result: list[Optional[str]] = [None] * n

    located = sorted([t for t in topics if t.char_start is not None], key=lambda t: t.char_start)
    if not located:
        return result

    order = sorted(range(n), key=lambda i: chunks_with_offsets[i][1])

    topic_idx = 0
    current_title: Optional[str] = None
    current_end: Optional[int] = None
    num_topics = len(located)

    for i in order:
        offset = chunks_with_offsets[i][1]
        while topic_idx < num_topics and located[topic_idx].char_start <= offset:
            current_title = located[topic_idx].title
            current_end = located[topic_idx].char_end
            topic_idx += 1
        if current_title is not None and (current_end is None or offset < current_end):
            result[i] = current_title

    return result